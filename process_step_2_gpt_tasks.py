from docx import Document
from docx.shared import Inches
from pathlib import Path
import os
import glob
from omegaconf import OmegaConf

from utils import read_task_responses

'''
Just a script to generate a docx file to dump generated tasks and plots to verify them.
'''

config_path = "configs/config.yaml"
config = OmegaConf.load(config_path)
openai_token_file = config.openai_token_file

dataset_folder = Path(config.dataset_valid_step_1)
response_path = dataset_folder / "gpt_tasks.jsonl"
# response_path_detailed = dataset_folder / "gpt_tasks_detailed.jsonl"

response = read_task_responses(response_path)
# response_detailed = read_task_responses(response_path_detailed)
dp_ids = sorted(list(response.keys()))

doc = Document()
# section = doc.sections[0]
# new_width, new_height = section.page_height, section.page_width
# section.page_width = new_width
# section.page_height = new_height

for idx in dp_ids:

    dp_folder = dataset_folder / str(idx)
    plot_files = glob.glob(os.path.join(str(dp_folder), "*.png"))
    plot_file = plot_files[0]

    table= doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'

    cell = table.cell(0, 0)
    cell.text = response[idx]

    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(plot_file, width=Inches(4.5))

    # cell = table.cell(0, 1)
    # cell.text = response_detailed[idx]

    doc.add_page_break()

doc.save('out/tasks.docx')


pass
pass
pass