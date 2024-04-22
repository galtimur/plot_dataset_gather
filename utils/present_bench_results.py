import base64
import glob
import json
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from omegaconf import OmegaConf

parent_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
sys.path.append(parent_dir)
print(parent_dir)
from utils import read_responses


def decode_image(encoded_image, output_image_file):
    decoded_image = base64.b64decode(encoded_image)

    # Write the decoded image data to a file
    with open(output_image_file, "wb") as image_file:
        image_file.write(decoded_image)


"""
Just a script to generate a docx file to dump generated tasks and plots to verify them.
"""
do_random = False
# suffix = "_probs"
suffix = ""
if do_random and len(suffix) == 0:
    suffix = "_random"

# config_path = "../configs/config.yaml"
config_path = "configs/config.yaml"
config = OmegaConf.load(config_path)

dataset_folder = Path(config.dataset_final)
results_folder = Path(config.out_folder)
temp_folder = results_folder / "temp"
os.makedirs(temp_folder, exist_ok=True)
bench_file = results_folder / f"benchmark_results.jsonl"
response_file = results_folder / "gpt_plots_dev.jsonl"

bench_scores = read_responses(bench_file)
plot_responses = read_responses(response_file)

temp_image_file = temp_folder / "plot.png"

# list of strings
ids = list(bench_scores.keys())

doc = Document()
section = doc.sections[0]
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height

for idx in ids:
    response = plot_responses[idx]
    result = bench_scores[idx]

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.add_run(f"ID = {idx}\n")
    paragraph.add_run(f'Score = {result["score"]:.0f}\n')
    if len(result["error"]) > 0 and not do_random:
        paragraph.add_run(f'Error = {result["error"]}\n')

    if not do_random:
        dp_folder = dataset_folder / str(idx)
    else:
        if "id_rnd" in result:
            rnd_idx = result["id_rnd"]
        else:
            rnd_idx = idx

        dp_folder = dataset_folder / str(rnd_idx)
        paragraph.add_run(f"RANDOM PAIR\n")

    plot_files = glob.glob(os.path.join(str(dp_folder), "*.png"))
    plot_file = plot_files[0]

    if len(plot_files) > 1 and not do_random:
        paragraph.add_run(
            f"There should be {len(plot_files)} images in GT, used only one\n"
        )

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    cell = table.cell(0, 0)
    cell.text = "Generated"

    if len(response["plot results"]["images"]) > 0:
        decode_image(response["plot results"]["images"][0], temp_image_file)

        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(str(temp_image_file), width=Inches(4))

    cell = table.cell(0, 1)
    cell.text = "Ground truth"
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(plot_file, width=Inches(4))

    doc.add_page_break()

doc.save("out/bench_results.docx")
